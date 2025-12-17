#!/usr/bin/env python3
"""
Script untuk menambahkan perhitungan manual TomBERT ke BAB 5 Pembahasan
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_tombert_manual_calculations(doc_path):
    """Menambahkan perhitungan manual TomBERT ke dokumen BAB 5"""
    
    doc = Document(doc_path)
    
    # Cari section 5.4 atau buat baru
    section_found = False
    insert_pos = None
    
    for i, para in enumerate(doc.paragraphs):
        if "5.4 Perhitungan Manual untuk Optimasi" in para.text or "5.4" in para.text and "Perhitungan Manual" in para.text:
            section_found = True
            insert_pos = i
            break
    
    # Jika tidak ditemukan, cari section terakhir
    if not section_found:
        for i, para in enumerate(doc.paragraphs):
            if "5.3" in para.text and "Implementasi" in para.text:
                insert_pos = i + 50  # Insert setelah section 5.3
                break
    
    if insert_pos is None:
        insert_pos = len(doc.paragraphs)
    
    # Tambahkan section baru untuk perhitungan manual TomBERT
    new_section = doc.paragraphs[insert_pos].insert_paragraph_before()
    new_section.add_run("5.4 Perhitungan Manual TomBERT").bold = True
    new_section.style = 'Heading 1'
    
    # 5.4.1 Forward Pass Calculation
    p = doc.paragraphs[insert_pos + 1].insert_paragraph_before()
    p.add_run("5.4.1 Perhitungan Forward Pass TomBERT").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 2].insert_paragraph_before()
    p.add_run("Forward pass pada TomBERT melibatkan beberapa tahap transformasi dari input teks dan gambar hingga menghasilkan prediksi sentimen. Berikut adalah perhitungan manual untuk setiap tahap:")
    
    # Step 1: Text Encoding
    p = doc.paragraphs[insert_pos + 3].insert_paragraph_before()
    p.add_run("Langkah 1: Encoding Teks dengan BERT").bold = True
    
    p = doc.paragraphs[insert_pos + 4].insert_paragraph_before()
    p.add_run("Input teks di-tokenize menjadi sequence of token IDs. Misalkan input teks: \"I love this product\" dengan target entity \"product\". ")
    p.add_run("Setelah tokenization dan padding, diperoleh:")
    
    p = doc.paragraphs[insert_pos + 5].insert_paragraph_before()
    p.add_run("• input_ids (full text): [101, 1045, 2293, 2023, 4284, 102] (length = 6)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 6].insert_paragraph_before()
    p.add_run("• s2_input_ids (entity): [101, 4284, 102] (length = 3)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 7].insert_paragraph_before()
    p.add_run("BERT encoder memproses input_ids melalui 12 layer transformer:")
    
    p = doc.paragraphs[insert_pos + 8].insert_paragraph_before()
    p.add_run("H^s = BERT(input_ids) ∈ ℝ^(batch_size × seq_len × hidden_size)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 9].insert_paragraph_before()
    p.add_run("Dengan hidden_size = 768, untuk batch_size = 1 dan seq_len = 6:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 10].insert_paragraph_before()
    p.add_run("H^s ∈ ℝ^(1 × 6 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 11].insert_paragraph_before()
    p.add_run("Pooled output (CLS token): pooled_output ∈ ℝ^(1 × 768)")
    p.style = 'List Bullet'
    
    # Step 2: Entity Encoding
    p = doc.paragraphs[insert_pos + 12].insert_paragraph_before()
    p.add_run("Langkah 2: Encoding Entity dengan s2_BERT").bold = True
    
    p = doc.paragraphs[insert_pos + 13].insert_paragraph_before()
    p.add_run("Entity \"product\" di-encode menggunakan BERT kedua (s2_bert):")
    
    p = doc.paragraphs[insert_pos + 14].insert_paragraph_before()
    p.add_run("H^e = s2_BERT(s2_input_ids) ∈ ℝ^(batch_size × s2_seq_len × hidden_size)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 15].insert_paragraph_before()
    p.add_run("Untuk s2_seq_len = 3: H^e ∈ ℝ^(1 × 3 × 768)")
    p.style = 'List Bullet'
    
    # Step 3: Image Feature Extraction
    p = doc.paragraphs[insert_pos + 16].insert_paragraph_before()
    p.add_run("Langkah 3: Ekstraksi Fitur Gambar dengan ResNet").bold = True
    
    p = doc.paragraphs[insert_pos + 17].insert_paragraph_before()
    p.add_run("Gambar di-resize menjadi 224×224 dan diproses oleh ResNet-152:")
    
    p = doc.paragraphs[insert_pos + 18].insert_paragraph_before()
    p.add_run("• ResNet menghasilkan feature map: 2048 × 7 × 7 = 2048 × 49")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 19].insert_paragraph_before()
    p.add_run("• visual_embeds_att ∈ ℝ^(batch_size × 2048 × 49)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 20].insert_paragraph_before()
    p.add_run("• Reshape: vis_embed_map = visual_embeds_att.view(-1, 2048, 49).permute(0, 2, 1)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 21].insert_paragraph_before()
    p.add_run("• Hasil: vis_embed_map ∈ ℝ^(1 × 49 × 2048)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 22].insert_paragraph_before()
    p.add_run("Proyeksi ke hidden dimension menggunakan linear layer:")
    
    p = doc.paragraphs[insert_pos + 23].insert_paragraph_before()
    p.add_run("converted_vis_embed_map = vismap2text(vis_embed_map)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 24].insert_paragraph_before()
    p.add_run("vismap2text: Linear(2048 → 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 25].insert_paragraph_before()
    p.add_run("converted_vis_embed_map ∈ ℝ^(1 × 49 × 768)")
    p.style = 'List Bullet'
    
    # Step 4: Cross-Attention
    p = doc.paragraphs[insert_pos + 26].insert_paragraph_before()
    p.add_run("Langkah 4: Cross-Attention antara Entity dan Image").bold = True
    
    p = doc.paragraphs[insert_pos + 27].insert_paragraph_before()
    p.add_run("Entity representation (H^e) digunakan sebagai query, image features sebagai key dan value:")
    
    p = doc.paragraphs[insert_pos + 28].insert_paragraph_before()
    p.add_run("s2_cross_encoder = ent2img_attention(H^e, converted_vis_embed_map, img_mask)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 29].insert_paragraph_before()
    p.add_run("Attention mechanism menghitung:")
    
    p = doc.paragraphs[insert_pos + 30].insert_paragraph_before()
    p.add_run("Q = H^e × W_q ∈ ℝ^(1 × 3 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 31].insert_paragraph_before()
    p.add_run("K = converted_vis_embed_map × W_k ∈ ℝ^(1 × 49 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 32].insert_paragraph_before()
    p.add_run("V = converted_vis_embed_map × W_v ∈ ℝ^(1 × 49 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 33].insert_paragraph_before()
    p.add_run("Attention scores: scores = (Q × K^T) / √d_k")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 34].insert_paragraph_before()
    p.add_run("dengan d_k = 768 / num_heads = 768 / 12 = 64")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 35].insert_paragraph_before()
    p.add_run("scores ∈ ℝ^(1 × 3 × 49)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 36].insert_paragraph_before()
    p.add_run("Apply mask dan softmax: attention_probs = softmax(scores + mask)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 37].insert_paragraph_before()
    p.add_run("Context: context = attention_probs × V ∈ ℝ^(1 × 3 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 38].insert_paragraph_before()
    p.add_run("Pooled output: s2_cross_output = pooler(context) ∈ ℝ^(1 × 768)")
    p.style = 'List Bullet'
    
    # Step 5: Multimodal Fusion
    p = doc.paragraphs[insert_pos + 39].insert_paragraph_before()
    p.add_run("Langkah 5: Fusi Multimodal").bold = True
    
    p = doc.paragraphs[insert_pos + 40].insert_paragraph_before()
    p.add_run("Image-aware entity representation digabung dengan text sequence:")
    
    p = doc.paragraphs[insert_pos + 41].insert_paragraph_before()
    p.add_run("transpose_img_embed = s2_cross_output.unsqueeze(1) ∈ ℝ^(1 × 1 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 42].insert_paragraph_before()
    p.add_run("text_img_output = concat([transpose_img_embed, H^s], dim=1)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 43].insert_paragraph_before()
    p.add_run("text_img_output ∈ ℝ^(1 × 7 × 768)  [1 image token + 6 text tokens]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 44].insert_paragraph_before()
    p.add_run("Multimodal encoder memproses kombinasi ini:")
    
    p = doc.paragraphs[insert_pos + 45].insert_paragraph_before()
    p.add_run("multimodal_encoder = comb_attention(text_img_output, attention_mask)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 46].insert_paragraph_before()
    p.add_run("img_att_text_output_layer ∈ ℝ^(1 × 7 × 768)")
    p.style = 'List Bullet'
    
    # Step 6: Pooling dan Classification
    p = doc.paragraphs[insert_pos + 47].insert_paragraph_before()
    p.add_run("Langkah 6: Pooling dan Klasifikasi").bold = True
    
    p = doc.paragraphs[insert_pos + 48].insert_paragraph_before()
    p.add_run("Untuk pooling=\"first\" (menggunakan image token):")
    
    p = doc.paragraphs[insert_pos + 49].insert_paragraph_before()
    p.add_run("comb_img_output = img_pooler(img_att_text_output_layer[:, 0, :])")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 50].insert_paragraph_before()
    p.add_run("pooled_output = dropout(comb_img_output) ∈ ℝ^(1 × 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 51].insert_paragraph_before()
    p.add_run("Logits: logits = classifier(pooled_output) ∈ ℝ^(1 × num_labels)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 52].insert_paragraph_before()
    p.add_run("Untuk 3 kelas sentimen: logits ∈ ℝ^(1 × 3)")
    p.style = 'List Bullet'
    
    # 5.4.2 Attention Mechanism Calculation
    p = doc.paragraphs[insert_pos + 53].insert_paragraph_before()
    p.add_run("5.4.2 Perhitungan Manual Attention Mechanism").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 54].insert_paragraph_before()
    p.add_run("Attention mechanism pada TomBERT menggunakan scaled dot-product attention. Berikut perhitungan detail:")
    
    p = doc.paragraphs[insert_pos + 55].insert_paragraph_before()
    p.add_run("Contoh perhitungan untuk 1 head attention (d_k = 64):")
    
    p = doc.paragraphs[insert_pos + 56].insert_paragraph_before()
    p.add_run("Misalkan:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 57].insert_paragraph_before()
    p.add_run("Q ∈ ℝ^(1 × 3 × 64)  [3 entity tokens, 64 dim per head]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 58].insert_paragraph_before()
    p.add_run("K ∈ ℝ^(1 × 49 × 64)  [49 image patches, 64 dim per head]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 59].insert_paragraph_before()
    p.add_run("V ∈ ℝ^(1 × 49 × 64)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 60].insert_paragraph_before()
    p.add_run("Step 1: Compute attention scores")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 61].insert_paragraph_before()
    p.add_run("scores = Q × K^T / √64")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 62].insert_paragraph_before()
    p.add_run("scores ∈ ℝ^(1 × 3 × 49)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 63].insert_paragraph_before()
    p.add_run("Contoh nilai untuk token pertama entity:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 64].insert_paragraph_before()
    p.add_run("scores[0, 0, :] = [2.3, 1.8, 0.5, -0.2, ..., 0.1]  [49 nilai]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 65].insert_paragraph_before()
    p.add_run("Step 2: Apply mask (untuk padding/image regions yang tidak valid)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 66].insert_paragraph_before()
    p.add_run("masked_scores = scores + mask")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 67].insert_paragraph_before()
    p.add_run("mask = -10000 untuk posisi yang di-mask, 0 untuk posisi valid")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 68].insert_paragraph_before()
    p.add_run("Step 3: Softmax normalization")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 69].insert_paragraph_before()
    p.add_run("attention_probs = softmax(masked_scores, dim=-1)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 70].insert_paragraph_before()
    p.add_run("attention_probs[0, 0, :] = [0.35, 0.28, 0.12, 0.08, ..., 0.02]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 71].insert_paragraph_before()
    p.add_run("Σ attention_probs[0, 0, :] = 1.0  [normalized]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 72].insert_paragraph_before()
    p.add_run("Step 4: Weighted sum of values")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 73].insert_paragraph_before()
    p.add_run("context = attention_probs × V")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 74].insert_paragraph_before()
    p.add_run("context[0, 0, :] = Σ(attention_probs[0, 0, i] × V[0, i, :]) untuk i = 0..48")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 75].insert_paragraph_before()
    p.add_run("context ∈ ℝ^(1 × 3 × 64)")
    p.style = 'List Bullet'
    
    # 5.4.3 Loss Calculation
    p = doc.paragraphs[insert_pos + 76].insert_paragraph_before()
    p.add_run("5.4.3 Perhitungan Manual Loss Function").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 77].insert_paragraph_before()
    p.add_run("TomBERT menggunakan Combined Loss yang menggabungkan Label Smoothing Cross-Entropy dan Focal Loss:")
    
    p = doc.paragraphs[insert_pos + 78].insert_paragraph_before()
    p.add_run("L_total = λ_ls × L_label_smoothing + λ_focal × L_focal")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 79].insert_paragraph_before()
    p.add_run("dengan λ_ls = 0.6 dan λ_focal = 0.4 (default)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 80].insert_paragraph_before()
    p.add_run("Contoh perhitungan untuk 1 sample:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 81].insert_paragraph_before()
    p.add_run("Logits dari model: logits = [2.1, 0.5, -1.2]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 82].insert_paragraph_before()
    p.add_run("True label: y_true = 0 (negative)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 83].insert_paragraph_before()
    p.add_run("Step 1: Softmax untuk mendapatkan probabilities")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 84].insert_paragraph_before()
    p.add_run("exp_logits = [exp(2.1), exp(0.5), exp(-1.2)] = [8.17, 1.65, 0.30]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 85].insert_paragraph_before()
    p.add_run("sum_exp = 8.17 + 1.65 + 0.30 = 10.12")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 86].insert_paragraph_before()
    p.add_run("p = [8.17/10.12, 1.65/10.12, 0.30/10.12] = [0.807, 0.163, 0.030]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 87].insert_paragraph_before()
    p.add_run("Step 2: Label Smoothing Cross-Entropy Loss")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 88].insert_paragraph_before()
    p.add_run("Label smoothing factor: α = 0.1 (adaptive, menurun selama training)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 89].insert_paragraph_before()
    p.add_run("Smooth label: y_smooth = [1-α, α/2, α/2] = [0.9, 0.05, 0.05]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 90].insert_paragraph_before()
    p.add_run("L_label_smoothing = -Σ(y_smooth[i] × log(p[i]))")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 91].insert_paragraph_before()
    p.add_run("L_label_smoothing = -(0.9×log(0.807) + 0.05×log(0.163) + 0.05×log(0.030))")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 92].insert_paragraph_before()
    p.add_run("L_label_smoothing = -(0.9×(-0.214) + 0.05×(-1.81) + 0.05×(-3.51))")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 93].insert_paragraph_before()
    p.add_run("L_label_smoothing = -(-0.193 - 0.091 - 0.176) = 0.460")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 94].insert_paragraph_before()
    p.add_run("Step 3: Focal Loss")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 95].insert_paragraph_before()
    p.add_run("Focal Loss parameters: α = 1.0, γ = 2.0")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 96].insert_paragraph_before()
    p.add_run("p_t = p[y_true] = p[0] = 0.807")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 97].insert_paragraph_before()
    p.add_run("CE_loss = -log(p_t) = -log(0.807) = 0.214")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 98].insert_paragraph_before()
    p.add_run("(1 - p_t)^γ = (1 - 0.807)^2 = (0.193)^2 = 0.037")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 99].insert_paragraph_before()
    p.add_run("L_focal = α × (1 - p_t)^γ × CE_loss = 1.0 × 0.037 × 0.214 = 0.008")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 100].insert_paragraph_before()
    p.add_run("Step 4: Combined Loss")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 101].insert_paragraph_before()
    p.add_run("L_total = 0.6 × 0.460 + 0.4 × 0.008 = 0.276 + 0.003 = 0.279")
    p.style = 'List Bullet'
    
    # 5.4.4 Gradient Calculation
    p = doc.paragraphs[insert_pos + 102].insert_paragraph_before()
    p.add_run("5.4.4 Perhitungan Manual Gradient (Backpropagation)").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 103].insert_paragraph_before()
    p.add_run("Backpropagation menghitung gradient loss terhadap semua parameter model:")
    
    p = doc.paragraphs[insert_pos + 104].insert_paragraph_before()
    p.add_run("∂L/∂logits = ∂L_total/∂logits")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 105].insert_paragraph_before()
    p.add_run("Untuk combined loss:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 106].insert_paragraph_before()
    p.add_run("∂L/∂logits = 0.6 × ∂L_label_smoothing/∂logits + 0.4 × ∂L_focal/∂logits")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 107].insert_paragraph_before()
    p.add_run("∂L/∂logits[i] = p[i] - y_smooth[i]  (untuk label smoothing)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 108].insert_paragraph_before()
    p.add_run("∂L/∂logits = [0.807-0.9, 0.163-0.05, 0.030-0.05] = [-0.093, 0.113, -0.020]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 109].insert_paragraph_before()
    p.add_run("Gradient kemudian di-backpropagate melalui classifier, pooling, attention layers, hingga ke BERT embeddings.")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 110].insert_paragraph_before()
    p.add_run("Gradient clipping diterapkan jika ||grad|| > max_grad_norm (default: 1.0)")
    p.style = 'List Bullet'
    
    # 5.4.5 Prediction Calculation
    p = doc.paragraphs[insert_pos + 111].insert_paragraph_before()
    p.add_run("5.4.5 Perhitungan Manual Prediksi").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 112].insert_paragraph_before()
    p.add_run("Setelah forward pass, logits di-convert menjadi probabilities dan prediksi:")
    
    p = doc.paragraphs[insert_pos + 113].insert_paragraph_before()
    p.add_run("Contoh: logits = [2.1, 0.5, -1.2]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 114].insert_paragraph_before()
    p.add_run("Step 1: Softmax")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 115].insert_paragraph_before()
    p.add_run("exp_logits = [exp(2.1), exp(0.5), exp(-1.2)] = [8.17, 1.65, 0.30]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 116].insert_paragraph_before()
    p.add_run("sum_exp = 10.12")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 117].insert_paragraph_before()
    p.add_run("probabilities = [8.17/10.12, 1.65/10.12, 0.30/10.12] = [0.807, 0.163, 0.030]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 118].insert_paragraph_before()
    p.add_run("Step 2: Prediksi")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 119].insert_paragraph_before()
    p.add_run("predicted_class = argmax(probabilities) = argmax([0.807, 0.163, 0.030]) = 0")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 120].insert_paragraph_before()
    p.add_run("confidence = max(probabilities) = 0.807 (80.7%)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 121].insert_paragraph_before()
    p.add_run("Interpretasi: Model memprediksi \"negative\" dengan confidence 80.7%")
    p.style = 'List Bullet'
    
    # Save document
    doc.save(doc_path)
    print(f"Perhitungan manual TomBERT telah ditambahkan ke {doc_path}")

if __name__ == "__main__":
    doc_path = "output/BAB5_Pembahasan.docx"
    add_tombert_manual_calculations(doc_path)

