#!/usr/bin/env python3
"""
Script untuk menambahkan perhitungan manual TomBERT menggunakan data aktual dari hasil eksperimen
"""

import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_tombert_manual_calculations_real_data(doc_path):
    """Menambahkan perhitungan manual TomBERT dengan data aktual"""
    
    # Load data aktual
    results_path = "output/tombert_ultra_optimized_20251129_225613/ultra_results.json"
    config_path = "output/tombert_ultra_optimized_20251129_225613/ultra_config.json"
    eval_path = "output/tombert_ultra_optimized_20251129_225613_20251129_225622/eval_results.txt"
    
    with open(results_path, 'r') as f:
        results = json.load(f)
    
    with open(config_path, 'r') as f:
        config = json.load(f)
    
    # Parse eval results
    eval_results = {}
    with open(eval_path, 'r') as f:
        for line in f:
            if '=' in line:
                key, value = line.strip().split(' = ')
                try:
                    eval_results[key] = float(value)
                except:
                    eval_results[key] = value
    
    doc = Document(doc_path)
    
    # Cari section 5.4 atau buat baru
    section_found = False
    insert_pos = None
    
    for i, para in enumerate(doc.paragraphs):
        if "5.4 Perhitungan Manual TomBERT" in para.text:
            section_found = True
            # Hapus section lama dan buat ulang
            j = i + 1
            while j < len(doc.paragraphs) and not (doc.paragraphs[j].text.startswith("5.5") or doc.paragraphs[j].text.startswith("5.") and "5.4" not in doc.paragraphs[j].text):
                doc.paragraphs[j]._element.getparent().remove(doc.paragraphs[j]._element)
            insert_pos = i
            break
    
    if not section_found:
        for i, para in enumerate(doc.paragraphs):
            if "5.3" in para.text and "Implementasi" in para.text:
                insert_pos = i + 50
                break
    
    if insert_pos is None:
        insert_pos = len(doc.paragraphs)
    
    # Data aktual dari eksperimen
    test_acc = results['final_test_accuracy']
    test_f1 = results['final_test_f1']
    test_prec = results['final_test_precision']
    test_rec = results['final_test_recall']
    eval_loss = eval_results.get('eval_loss', 0.7992)
    
    label_smoothing = config.get('label_smoothing', 0.2)
    focal_alpha = config.get('focal_alpha', 1.0)
    focal_gamma = config.get('focal_gamma', 2.0)
    learning_rate = config.get('learning_rate', 1e-5)
    batch_size = config.get('train_batch_size', 16)
    max_seq_length = config.get('max_seq_length', 128)
    max_entity_length = config.get('max_entity_length', 32)
    pooling = config.get('pooling', 'concat')
    hidden_size = 768  # BERT base
    
    # Tambahkan section baru
    if not section_found:
        new_section = doc.paragraphs[insert_pos].insert_paragraph_before()
        new_section.add_run("5.4 Perhitungan Manual TomBERT").bold = True
        new_section.style = 'Heading 1'
        insert_pos += 1
    
    # Informasi konfigurasi aktual
    p = doc.paragraphs[insert_pos].insert_paragraph_before()
    p.add_run("5.4.0 Konfigurasi Eksperimen Aktual").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 1].insert_paragraph_before()
    p.add_run("Perhitungan manual berikut menggunakan konfigurasi dan hasil aktual dari eksperimen TomBERT pada dataset Twitter15:")
    
    p = doc.paragraphs[insert_pos + 2].insert_paragraph_before()
    p.add_run(f"• Model: TomBERT dengan pooling='{pooling}'")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 3].insert_paragraph_before()
    p.add_run(f"• BERT: bert-base-uncased (hidden_size={hidden_size})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 4].insert_paragraph_before()
    p.add_run(f"• Max sequence length: {max_seq_length} tokens")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 5].insert_paragraph_before()
    p.add_run(f"• Max entity length: {max_entity_length} tokens")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 6].insert_paragraph_before()
    p.add_run(f"• Batch size: {batch_size}")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 7].insert_paragraph_before()
    p.add_run(f"• Learning rate: {learning_rate}")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 8].insert_paragraph_before()
    p.add_run(f"• Label smoothing: {label_smoothing}")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 9].insert_paragraph_before()
    p.add_run(f"• Focal loss: α={focal_alpha}, γ={focal_gamma}")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 10].insert_paragraph_before()
    p.add_run(f"• Hasil Test Set: Accuracy={test_acc:.4f} ({test_acc*100:.2f}%), F1={test_f1:.4f} ({test_f1*100:.2f}%), Precision={test_prec:.4f} ({test_prec*100:.2f}%), Recall={test_rec:.4f} ({test_rec*100:.2f}%)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 11].insert_paragraph_before()
    p.add_run(f"• Eval Loss: {eval_loss:.4f}")
    p.style = 'List Bullet'
    
    # 5.4.1 Forward Pass dengan dimensi aktual
    p = doc.paragraphs[insert_pos + 12].insert_paragraph_before()
    p.add_run("5.4.1 Perhitungan Forward Pass TomBERT (Data Aktual)").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 13].insert_paragraph_before()
    p.add_run("Berikut perhitungan forward pass menggunakan konfigurasi aktual dari eksperimen:")
    
    p = doc.paragraphs[insert_pos + 14].insert_paragraph_before()
    p.add_run("Langkah 1: Encoding Teks dengan BERT").bold = True
    
    p = doc.paragraphs[insert_pos + 15].insert_paragraph_before()
    p.add_run(f"Input teks di-tokenize dan di-pad hingga max_seq_length={max_seq_length}:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 16].insert_paragraph_before()
    p.add_run(f"H^s = BERT(input_ids) ∈ ℝ^(batch_size × {max_seq_length} × {hidden_size})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 17].insert_paragraph_before()
    p.add_run(f"Untuk batch_size={batch_size}: H^s ∈ ℝ^({batch_size} × {max_seq_length} × {hidden_size})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 18].insert_paragraph_before()
    p.add_run("Langkah 2: Encoding Entity dengan s2_BERT").bold = True
    
    p = doc.paragraphs[insert_pos + 19].insert_paragraph_before()
    p.add_run(f"Entity di-encode dengan max_entity_length={max_entity_length}:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 20].insert_paragraph_before()
    p.add_run(f"H^e = s2_BERT(s2_input_ids) ∈ ℝ^({batch_size} × {max_entity_length} × {hidden_size})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 21].insert_paragraph_before()
    p.add_run("Langkah 3: Ekstraksi Fitur Gambar").bold = True
    
    p = doc.paragraphs[insert_pos + 22].insert_paragraph_before()
    p.add_run("ResNet-152 menghasilkan feature map 2048 × 7 × 7 = 2048 × 49:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 23].insert_paragraph_before()
    p.add_run("vis_embed_map ∈ ℝ^({batch_size} × 49 × 2048)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 24].insert_paragraph_before()
    p.add_run("Proyeksi: converted_vis_embed_map = Linear(2048 → 768)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 25].insert_paragraph_before()
    p.add_run("converted_vis_embed_map ∈ ℝ^({batch_size} × 49 × {hidden_size})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 26].insert_paragraph_before()
    p.add_run("Langkah 4: Cross-Attention Entity-Image").bold = True
    
    p = doc.paragraphs[insert_pos + 27].insert_paragraph_before()
    p.add_run("s2_cross_output ∈ ℝ^({batch_size} × {hidden_size})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 28].insert_paragraph_before()
    p.add_run("Langkah 5: Fusi Multimodal").bold = True
    
    p = doc.paragraphs[insert_pos + 29].insert_paragraph_before()
    p.add_run(f"Dengan pooling='{pooling}':")
    p.style = 'List Bullet'
    
    if pooling == 'concat':
        p = doc.paragraphs[insert_pos + 30].insert_paragraph_before()
        p.add_run("pooled_output = concat([text_pooler(output), img_pooler(output)]) ∈ ℝ^({batch_size} × {hidden_size*2})")
        p.style = 'List Bullet'
        
        p = doc.paragraphs[insert_pos + 31].insert_paragraph_before()
        p.add_run("logits = classifier(pooled_output) ∈ ℝ^({batch_size} × 3)")
        p.style = 'List Bullet'
    else:
        p = doc.paragraphs[insert_pos + 30].insert_paragraph_before()
        p.add_run(f"pooled_output ∈ ℝ^({batch_size} × {hidden_size})")
        p.style = 'List Bullet'
        
        p = doc.paragraphs[insert_pos + 31].insert_paragraph_before()
        p.add_run("logits = classifier(pooled_output) ∈ ℝ^({batch_size} × 3)")
        p.style = 'List Bullet'
    
    # 5.4.2 Loss Calculation dengan nilai aktual
    p = doc.paragraphs[insert_pos + 32].insert_paragraph_before()
    p.add_run("5.4.2 Perhitungan Manual Loss Function (Data Aktual)").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 33].insert_paragraph_before()
    p.add_run(f"TomBERT menggunakan Combined Loss dengan konfigurasi aktual:")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 34].insert_paragraph_before()
    p.add_run(f"L_total = 0.6 × L_label_smoothing + 0.4 × L_focal")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 35].insert_paragraph_before()
    p.add_run(f"dengan label_smoothing={label_smoothing}, focal_alpha={focal_alpha}, focal_gamma={focal_gamma}")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 36].insert_paragraph_before()
    p.add_run("Contoh perhitungan untuk 1 sample dari test set:")
    p.style = 'List Bullet'
    
    # Contoh numerik berdasarkan hasil aktual
    # Misalkan logits yang menghasilkan prediksi dengan confidence tinggi
    p = doc.paragraphs[insert_pos + 37].insert_paragraph_before()
    p.add_run("Logits dari model (contoh): logits = [2.5, 0.3, -1.8]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 38].insert_paragraph_before()
    p.add_run("True label: y_true = 0 (negative)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 39].insert_paragraph_before()
    p.add_run("Step 1: Softmax probabilities")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 40].insert_paragraph_before()
    p.add_run("exp_logits = [exp(2.5), exp(0.3), exp(-1.8)] = [12.18, 1.35, 0.17]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 41].insert_paragraph_before()
    p.add_run("sum_exp = 12.18 + 1.35 + 0.17 = 13.70")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 42].insert_paragraph_before()
    p.add_run("p = [12.18/13.70, 1.35/13.70, 0.17/13.70] = [0.889, 0.099, 0.012]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 43].insert_paragraph_before()
    p.add_run(f"Step 2: Label Smoothing Cross-Entropy (α={label_smoothing})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 44].insert_paragraph_before()
    p.add_run(f"y_smooth = [1-{label_smoothing}, {label_smoothing}/2, {label_smoothing}/2] = [0.8, 0.1, 0.1]")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 45].insert_paragraph_before()
    p.add_run("L_label_smoothing = -(0.8×log(0.889) + 0.1×log(0.099) + 0.1×log(0.012))")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 46].insert_paragraph_before()
    p.add_run("L_label_smoothing = -(0.8×(-0.118) + 0.1×(-2.313) + 0.1×(-4.423))")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 47].insert_paragraph_before()
    p.add_run("L_label_smoothing = -(-0.094 - 0.231 - 0.442) = 0.767")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 48].insert_paragraph_before()
    p.add_run(f"Step 3: Focal Loss (α={focal_alpha}, γ={focal_gamma})")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 49].insert_paragraph_before()
    p.add_run("p_t = p[0] = 0.889")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 50].insert_paragraph_before()
    p.add_run("CE_loss = -log(0.889) = 0.118")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 51].insert_paragraph_before()
    p.add_run(f"(1 - p_t)^γ = (1 - 0.889)^{focal_gamma} = (0.111)^{focal_gamma} = 0.012")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 52].insert_paragraph_before()
    p.add_run(f"L_focal = {focal_alpha} × 0.012 × 0.118 = 0.001")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 53].insert_paragraph_before()
    p.add_run("Step 4: Combined Loss")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 54].insert_paragraph_before()
    p.add_run("L_total = 0.6 × 0.767 + 0.4 × 0.001 = 0.460 + 0.0004 = 0.460")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 55].insert_paragraph_before()
    p.add_run(f"Rata-rata loss pada test set: {eval_loss:.4f} (sesuai dengan eval_results.txt)")
    p.style = 'List Bullet'
    
    # 5.4.3 Perhitungan metrik evaluasi aktual
    p = doc.paragraphs[insert_pos + 56].insert_paragraph_before()
    p.add_run("5.4.3 Perhitungan Manual Metrik Evaluasi (Data Aktual)").bold = True
    p.style = 'Heading 2'
    
    p = doc.paragraphs[insert_pos + 57].insert_paragraph_before()
    p.add_run("Berikut perhitungan manual metrik evaluasi berdasarkan hasil aktual pada test set (1037 samples):")
    
    p = doc.paragraphs[insert_pos + 58].insert_paragraph_before()
    p.add_run("Accuracy").bold = True
    
    p = doc.paragraphs[insert_pos + 59].insert_paragraph_before()
    p.add_run(f"Accuracy = (TP + TN) / Total = {test_acc:.4f} = {test_acc*100:.2f}%")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 60].insert_paragraph_before()
    p.add_run(f"Jumlah prediksi benar = {test_acc * 1037:.0f} dari 1037 samples")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 61].insert_paragraph_before()
    p.add_run("Macro Precision").bold = True
    
    p = doc.paragraphs[insert_pos + 62].insert_paragraph_before()
    p.add_run(f"Macro Precision = {test_prec:.4f} = {test_prec*100:.2f}%")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 63].insert_paragraph_before()
    p.add_run("Macro Recall").bold = True
    
    p = doc.paragraphs[insert_pos + 64].insert_paragraph_before()
    p.add_run(f"Macro Recall = {test_rec:.4f} = {test_rec*100:.2f}%")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 65].insert_paragraph_before()
    p.add_run("Macro F1-Score").bold = True
    
    p = doc.paragraphs[insert_pos + 66].insert_paragraph_before()
    p.add_run(f"Macro F1 = 2 × (Precision × Recall) / (Precision + Recall)")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 67].insert_paragraph_before()
    p.add_run(f"Macro F1 = 2 × ({test_prec:.4f} × {test_rec:.4f}) / ({test_prec:.4f} + {test_rec:.4f})")
    p.style = 'List Bullet'
    
    f1_calc = 2 * (test_prec * test_rec) / (test_prec + test_rec)
    p = doc.paragraphs[insert_pos + 68].insert_paragraph_before()
    p.add_run(f"Macro F1 = 2 × {test_prec*test_rec:.4f} / {test_prec+test_rec:.4f} = {f1_calc:.4f}")
    p.style = 'List Bullet'
    
    p = doc.paragraphs[insert_pos + 69].insert_paragraph_before()
    p.add_run(f"Nilai aktual dari model: {test_f1:.4f} ({test_f1*100:.2f}%)")
    p.style = 'List Bullet'
    
    # Save document
    doc.save(doc_path)
    print(f"Perhitungan manual TomBERT dengan data aktual telah ditambahkan ke {doc_path}")
    print(f"Data aktual yang digunakan:")
    print(f"  - Test Accuracy: {test_acc:.4f}")
    print(f"  - Test F1: {test_f1:.4f}")
    print(f"  - Test Precision: {test_prec:.4f}")
    print(f"  - Test Recall: {test_rec:.4f}")
    print(f"  - Eval Loss: {eval_loss:.4f}")

if __name__ == "__main__":
    doc_path = "output/BAB5_Pembahasan.docx"
    add_tombert_manual_calculations_real_data(doc_path)

